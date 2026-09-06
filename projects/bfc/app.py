#!/usr/bin/env python3
"""
Document Utilities
===================
A single-file Flask app for converting and splitting documents (PDF, DOCX,
EPUB, Markdown, HTML, TXT), with optional AI-assisted chapter splitting and
table-of-contents generation (via a user-supplied Claude or Gemini API key).

Run:
    pip install -r requirements.txt
    python app.py

Then open http://127.0.0.1:5000

See README.md for configuration, limitations, and security notes.
"""

import os
import re
import io
import json
import time
import uuid
import html
import shutil
import zipfile
import tempfile
import threading
import traceback
from xml.sax.saxutils import escape as xml_escape

from flask import Flask, request, jsonify, send_file, session, Response, abort
from werkzeug.utils import secure_filename

import requests

# ---------------------------------------------------------------------------
# Optional heavy dependencies. Each is imported defensively so the app still
# starts (and other formats still work) if one library isn't installed.
# ---------------------------------------------------------------------------
try:
    import PyMuPDF
except ImportError:
    fitz = None

try:
    from pypdf import PdfReader, PdfWriter
except ImportError:
    try:
        from PyPDF2 import PdfReader, PdfWriter  # legacy fallback name
    except ImportError:
        PdfReader = PdfWriter = None

try:
    import docx as docx_lib
    from docx import Document as DocxDocument
except ImportError:
    docx_lib = None
    DocxDocument = None

try:
    from ebooklib import epub
    import ebooklib
except ImportError:
    epub = None
    ebooklib = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
except ImportError:
    SimpleDocTemplate = None


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
APP_ROOT = tempfile.gettempdir()
STORAGE_DIR = os.path.join(APP_ROOT, "docutil_storage")
os.makedirs(STORAGE_DIR, exist_ok=True)

MAX_CONTENT_LENGTH = 60 * 1024 * 1024  # 60 MB upload cap
SESSION_TTL_SECONDS = 2 * 60 * 60      # temp files older than this get swept

ALLOWED_EXTENSIONS = {"pdf", "docx", "epub", "txt", "md", "markdown", "html", "htm"}
TARGET_FORMATS = {"pdf", "docx", "epub", "md", "html", "txt"}

DEFAULT_MODELS = {
    "claude": "claude-sonnet-5",
    "gemini": "gemini-3.1-flash-lite",
}
ENV_KEYS = {
    "claude": "ANTHROPIC_API_KEY",
    "gemini": "GEMINI_API_KEY",
}

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = MAX_CONTENT_LENGTH
app.secret_key = os.environ.get("DOCUTIL_SECRET_KEY") or os.urandom(32)

_ID_RE = re.compile(r"^[0-9a-f]{32}$")


# ---------------------------------------------------------------------------
# Session / storage helpers
# ---------------------------------------------------------------------------
def get_sid():
    sid = session.get("sid")
    if not sid:
        sid = uuid.uuid4().hex
        session["sid"] = sid
    return sid


def session_dir():
    d = os.path.join(STORAGE_DIR, get_sid())
    os.makedirs(d, exist_ok=True)
    return d


def valid_id(s):
    return bool(s) and bool(_ID_RE.match(s))


def safe_stem(filename):
    stem = filename.rsplit(".", 1)[0] if "." in filename else filename
    stem = secure_filename(stem) or "document"
    return stem[:80]


def sanitize_title(title):
    t = secure_filename((title or "").strip().replace(" ", "_"))[:60]
    return t or f"part_{uuid.uuid4().hex[:6]}"


_manifest_locks = {}
_manifest_locks_guard = threading.Lock()


def _get_lock(sid):
    with _manifest_locks_guard:
        if sid not in _manifest_locks:
            _manifest_locks[sid] = threading.Lock()
        return _manifest_locks[sid]


def _manifest_path():
    return os.path.join(session_dir(), "manifest.json")


def load_manifest():
    p = _manifest_path()
    if not os.path.exists(p):
        return {"files": {}, "outputs": {}}
    try:
        with open(p, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {"files": {}, "outputs": {}}


def save_manifest(m):
    p = _manifest_path()
    tmp = p + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(m, f)
    os.replace(tmp, p)


def register_file(file_id, filename, ext, stored_name, meta):
    sid = get_sid()
    with _get_lock(sid):
        m = load_manifest()
        m["files"][file_id] = {
            "filename": filename, "ext": ext, "stored_name": stored_name,
            "meta": meta, "created": time.time(),
        }
        save_manifest(m)


def register_output(output_id, filename, ext, stored_name):
    sid = get_sid()
    with _get_lock(sid):
        m = load_manifest()
        m["outputs"][output_id] = {
            "filename": filename, "ext": ext, "stored_name": stored_name,
            "created": time.time(),
        }
        save_manifest(m)


def get_file_record(file_id):
    m = load_manifest()
    rec = m["files"].get(file_id)
    if not rec:
        return None
    rec = dict(rec)
    rec["abs_path"] = os.path.join(session_dir(), rec["stored_name"])
    if not os.path.exists(rec["abs_path"]):
        return None
    return rec


def get_output_record(output_id):
    m = load_manifest()
    rec = m["outputs"].get(output_id)
    if not rec:
        return None
    rec = dict(rec)
    rec["abs_path"] = os.path.join(session_dir(), rec["stored_name"])
    if not os.path.exists(rec["abs_path"]):
        return None
    return rec


def cleanup_loop():
    while True:
        try:
            now = time.time()
            if os.path.isdir(STORAGE_DIR):
                for name in os.listdir(STORAGE_DIR):
                    p = os.path.join(STORAGE_DIR, name)
                    try:
                        if os.path.isdir(p) and (now - os.path.getmtime(p)) > SESSION_TTL_SECONDS:
                            shutil.rmtree(p, ignore_errors=True)
                    except Exception:
                        pass
        except Exception:
            pass
        time.sleep(900)


threading.Thread(target=cleanup_loop, daemon=True).start()


# ---------------------------------------------------------------------------
# Document model: every source format is parsed into a flat list of "blocks"
#   {'type': 'heading', 'level': 1..6, 'text': str, 'page'/'chapter': int}
#   {'type': 'para', 'text': str, 'page'/'chapter': int}
# which every target renderer consumes. This keeps the N-format conversion
# matrix to N extractors + N renderers instead of N^2 converters. Fidelity is
# intentionally text/structure-level (headings + paragraphs) — images, exact
# layout, and typography are not preserved. See README for details.
# ---------------------------------------------------------------------------

# ---- Extractors -------------------------------------------------------

def extract_blocks(path, ext):
    ext = ext.lower()
    if ext == "pdf":
        return extract_pdf(path)
    if ext == "docx":
        return extract_docx(path)
    if ext == "epub":
        return extract_epub(path)
    if ext == "txt":
        return extract_txt(path)
    if ext in ("md", "markdown"):
        return extract_md(path)
    if ext in ("html", "htm"):
        return extract_html(path)
    raise ValueError(f"Unsupported source format: .{ext}")


def extract_pdf(path):
    if fitz is not None:
        return _extract_pdf_fitz(path)
    if PdfReader is not None:
        return _extract_pdf_pypdf(path)
    raise RuntimeError(
        "Reading PDF files requires PyMuPDF or pypdf. Install with: "
        "pip install PyMuPDF pypdf"
    )


def _extract_pdf_fitz(path):
    doc = fitz.open(path)
    sizes = []
    raw_blocks = []
    for pno in range(len(doc)):
        page = doc[pno]
        d = page.get_text("dict")
        for block in d.get("blocks", []):
            if block.get("type") != 0:
                continue
            parts = []
            max_size = 0.0
            bold = False
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    t = span.get("text", "")
                    if t.strip():
                        parts.append(t)
                    max_size = max(max_size, span.get("size", 0) or 0)
                    if "bold" in (span.get("font", "") or "").lower():
                        bold = True
            text = re.sub(r"\s+", " ", " ".join(parts)).strip()
            if not text:
                continue
            sizes.append(max_size)
            raw_blocks.append({"page": pno, "text": text, "size": max_size, "bold": bold})
    doc.close()
    if not sizes:
        return []
    sizes_sorted = sorted(sizes)
    body_size = sizes_sorted[len(sizes_sorted) // 2] or 1.0
    blocks = []
    for b in raw_blocks:
        ratio = b["size"] / body_size
        if ratio >= 1.45:
            level = 1
        elif ratio >= 1.2:
            level = 2
        elif ratio >= 1.08 or (b["bold"] and len(b["text"]) < 90):
            level = 3
        else:
            level = None
        if level:
            blocks.append({"type": "heading", "level": level, "text": b["text"], "page": b["page"]})
        else:
            blocks.append({"type": "para", "text": b["text"], "page": b["page"]})
    return blocks


def _extract_pdf_pypdf(path):
    reader = PdfReader(path)
    blocks = []
    for pno, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        for para in re.split(r"\n\s*\n", text):
            para = re.sub(r"\s+", " ", para).strip()
            if para:
                blocks.append({"type": "para", "text": para, "page": pno})
    return blocks


def extract_docx(path):
    if DocxDocument is None:
        raise RuntimeError("python-docx is required to read DOCX files. Install with: pip install python-docx")
    d = DocxDocument(path)
    blocks = []
    for p in d.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = (p.style.name or "") if p.style else ""
        m = re.match(r"Heading\s*(\d+)", style, re.I)
        if m:
            blocks.append({"type": "heading", "level": min(int(m.group(1)), 6), "text": text})
        elif style.lower() == "title":
            blocks.append({"type": "heading", "level": 1, "text": text})
        else:
            blocks.append({"type": "para", "text": text})
    return blocks


def extract_epub(path):
    if epub is None or BeautifulSoup is None:
        raise RuntimeError(
            "EbookLib and beautifulsoup4 are required to read EPUB files. "
            "Install with: pip install EbookLib beautifulsoup4 lxml"
        )
    book = epub.read_epub(path)
    spine_ids = [i[0] for i in book.spine if i[0] != "nav"]
    id_to_item = {item.get_id(): item for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT)}
    blocks = []
    for chapter_idx, item_id in enumerate(spine_ids):
        item = id_to_item.get(item_id)
        if item is None:
            continue
        try:
            soup = BeautifulSoup(item.get_content(), "lxml")
        except Exception:
            soup = BeautifulSoup(item.get_content(), "html.parser")
        for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p"]):
            text = re.sub(r"\s+", " ", tag.get_text(strip=True))
            if not text:
                continue
            if tag.name.startswith("h"):
                blocks.append({"type": "heading", "level": int(tag.name[1]), "text": text, "chapter": chapter_idx})
            else:
                blocks.append({"type": "para", "text": text, "chapter": chapter_idx})
    return blocks


def extract_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()
    blocks = []
    for para in re.split(r"\n\s*\n", content):
        para = para.strip()
        if not para:
            continue
        single_line = "\n" not in para
        if single_line and (para.isupper() and len(para) < 100):
            blocks.append({"type": "heading", "level": 1, "text": para.strip()})
        else:
            blocks.append({"type": "para", "text": re.sub(r"\s+", " ", para)})
    return blocks


def extract_md(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()
    blocks = []
    buf = []

    def flush():
        if buf:
            text = re.sub(r"\s+", " ", " ".join(buf)).strip()
            if text:
                blocks.append({"type": "para", "text": text})
            buf.clear()

    for line in content.splitlines():
        stripped = line.strip()
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            flush()
            level = len(m.group(1))
            text = m.group(2).strip()
            if text:
                blocks.append({"type": "heading", "level": level, "text": text})
        elif not stripped:
            flush()
        else:
            buf.append(stripped)
    flush()
    return blocks


def extract_html(path):
    if BeautifulSoup is None:
        raise RuntimeError("beautifulsoup4 is required to read HTML files. Install with: pip install beautifulsoup4 lxml")
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()
    try:
        soup = BeautifulSoup(content, "lxml")
    except Exception:
        soup = BeautifulSoup(content, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    blocks = []
    body = soup.body or soup
    for tag in body.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p"]):
        text = re.sub(r"\s+", " ", tag.get_text(strip=True))
        if not text:
            continue
        if tag.name.startswith("h"):
            blocks.append({"type": "heading", "level": int(tag.name[1]), "text": text})
        else:
            blocks.append({"type": "para", "text": text})
    return blocks


# ---- Renderers ----------------------------------------------------------

def render_blocks(blocks, ext, path):
    ext = ext.lower()
    if ext == "pdf":
        return render_pdf(blocks, path)
    if ext == "docx":
        return render_docx(blocks, path)
    if ext == "epub":
        return render_epub(blocks, path)
    if ext in ("md", "markdown"):
        return render_md(blocks, path)
    if ext in ("html", "htm"):
        return render_html(blocks, path)
    if ext == "txt":
        return render_txt(blocks, path)
    raise ValueError(f"Unsupported target format: .{ext}")


def render_txt(blocks, path):
    out = []
    for b in blocks:
        if b["type"] == "heading":
            text = b["text"]
            out.append("")
            if b["level"] == 1:
                out.append(text.upper())
                out.append("=" * max(3, len(text)))
            else:
                out.append(text)
                out.append("-" * max(3, len(text)))
            out.append("")
        else:
            out.append(b["text"])
            out.append("")
    content = "\n".join(out).strip() + "\n"
    with open(path, "w", encoding="utf-8") as f:
        f.write(content if content.strip() else "(empty document)\n")


def render_md(blocks, path):
    out = []
    for b in blocks:
        if b["type"] == "heading":
            out.append("#" * min(max(b["level"], 1), 6) + " " + b["text"])
            out.append("")
        else:
            out.append(b["text"])
            out.append("")
    content = "\n".join(out).strip() + "\n"
    with open(path, "w", encoding="utf-8") as f:
        f.write(content if content.strip() else "*(empty document)*\n")


def render_html(blocks, path, title="Converted Document"):
    parts = [
        "<!DOCTYPE html>",
        "<html lang='en'><head><meta charset='utf-8'>",
        f"<title>{html.escape(title)}</title>",
        "<style>body{font-family:Georgia,'Iowan Old Style',serif;max-width:760px;margin:40px auto;"
        "padding:0 20px;line-height:1.65;color:#222}"
        "h1,h2,h3,h4,h5,h6{font-family:-apple-system,'Segoe UI',sans-serif;color:#111;"
        "margin-top:1.4em}</style></head><body>",
    ]
    any_content = False
    for b in blocks:
        any_content = True
        if b["type"] == "heading":
            lvl = min(max(b["level"], 1), 6)
            parts.append(f"<h{lvl}>{html.escape(b['text'])}</h{lvl}>")
        else:
            parts.append(f"<p>{html.escape(b['text'])}</p>")
    if not any_content:
        parts.append("<p><em>(empty document)</em></p>")
    parts.append("</body></html>")
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(parts))


def render_docx(blocks, path):
    if docx_lib is None or DocxDocument is None:
        raise RuntimeError("python-docx is required to create DOCX files. Install with: pip install python-docx")
    d = DocxDocument()
    if not blocks:
        d.add_paragraph("(empty document)")
    for b in blocks:
        if b["type"] == "heading":
            d.add_heading(b["text"], level=min(max(b["level"], 1), 9))
        else:
            d.add_paragraph(b["text"])
    d.save(path)


def render_pdf(blocks, path):
    if SimpleDocTemplate is None:
        raise RuntimeError("reportlab is required to create PDF files. Install with: pip install reportlab")
    doc = SimpleDocTemplate(
        path, pagesize=LETTER,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
        topMargin=0.9 * inch, bottomMargin=0.9 * inch,
    )
    styles = getSampleStyleSheet()
    h_styles = {
        1: ParagraphStyle("H1", parent=styles["Heading1"], spaceBefore=18, spaceAfter=8),
        2: ParagraphStyle("H2", parent=styles["Heading2"], spaceBefore=14, spaceAfter=6),
        3: ParagraphStyle("H3", parent=styles["Heading3"], spaceBefore=10, spaceAfter=4),
    }
    body_style = styles["BodyText"]
    story = []
    for b in blocks:
        text = xml_escape(b["text"])
        if b["type"] == "heading":
            lvl = min(b["level"], 3)
            story.append(Paragraph(text, h_styles.get(lvl, h_styles[3])))
        else:
            story.append(Paragraph(text, body_style))
            story.append(Spacer(1, 6))
    if not story:
        story.append(Paragraph("(empty document)", body_style))
    doc.build(story)


def render_epub(blocks, path, book_title="Converted Document"):
    if epub is None:
        raise RuntimeError("EbookLib is required to create EPUB files. Install with: pip install EbookLib")
    book = epub.EpubBook()
    book.set_identifier(str(uuid.uuid4()))
    book.set_title(book_title)
    book.set_language("en")

    chapters_blocks = []
    current = []
    for b in blocks:
        if b["type"] == "heading" and b["level"] == 1 and current:
            chapters_blocks.append(current)
            current = [b]
        else:
            current.append(b)
    if current:
        chapters_blocks.append(current)
    if not chapters_blocks:
        chapters_blocks = [[{"type": "para", "text": "(empty document)"}]]

    epub_items = []
    for i, chap_blocks in enumerate(chapters_blocks):
        title = next((b["text"] for b in chap_blocks if b["type"] == "heading"), f"Chapter {i + 1}")
        html_parts = []
        for b in chap_blocks:
            safe = html.escape(b["text"])
            if b["type"] == "heading":
                lvl = min(max(b["level"], 1), 6)
                html_parts.append(f"<h{lvl}>{safe}</h{lvl}>")
            else:
                html_parts.append(f"<p>{safe}</p>")
        content = f"<html><head><title>{html.escape(title)}</title></head><body>{''.join(html_parts)}</body></html>"
        c = epub.EpubHtml(title=title, file_name=f"chap_{i + 1}.xhtml", lang="en")
        c.set_content(content)
        book.add_item(c)
        epub_items.append(c)

    book.toc = epub_items
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav"] + epub_items
    epub.write_epub(path, book)


# ---- Splitting ------------------------------------------------------------

def parse_ranges(spec, total):
    spec = (spec or "").strip()
    if not spec:
        raise ValueError("No range specified.")
    groups = []
    for token in spec.split(","):
        token = token.strip()
        if not token:
            continue
        if "-" in token:
            a_s, b_s = token.split("-", 1)
            a, b = int(a_s), int(b_s)
        else:
            a = b = int(token)
        if a < 1 or b < a or b > total:
            raise ValueError(f"Invalid range '{token}' for a document with {total} unit(s).")
        groups.append((a - 1, b - 1))
    if not groups:
        raise ValueError("No valid ranges found.")
    return groups


def group_blocks_into_sections(blocks):
    sections = []
    current = []
    for b in blocks:
        if b["type"] == "heading" and b["level"] == 1 and current:
            sections.append(current)
            current = [b]
        else:
            current.append(b)
    if current:
        sections.append(current)
    return sections


def _split_pdf(src_path, ranges, out_dir):
    if fitz is not None:
        src = fitz.open(src_path)
        outs = []
        for a, b in ranges:
            nd = fitz.open()
            nd.insert_pdf(src, from_page=a, to_page=b)
            p = os.path.join(out_dir, f"tmp_{uuid.uuid4().hex}.pdf")
            nd.save(p)
            nd.close()
            outs.append(p)
        src.close()
        return outs
    if PdfReader is not None and PdfWriter is not None:
        reader = PdfReader(src_path)
        outs = []
        for a, b in ranges:
            writer = PdfWriter()
            for i in range(a, b + 1):
                writer.add_page(reader.pages[i])
            p = os.path.join(out_dir, f"tmp_{uuid.uuid4().hex}.pdf")
            with open(p, "wb") as f:
                writer.write(f)
            outs.append(p)
        return outs
    raise RuntimeError("Splitting PDF files requires PyMuPDF or pypdf.")


def split_epub(src_path, ranges, out_dir):
    if epub is None:
        raise RuntimeError("EbookLib is required to split EPUB files.")
    book = epub.read_epub(src_path)
    spine_ids = [i[0] for i in book.spine if i[0] != "nav"]
    doc_items = {item.get_id(): item for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT)}
    resource_items = [item for item in book.get_items() if item.get_type() != ebooklib.ITEM_DOCUMENT]
    total = len(spine_ids)
    outs = []
    for a, b in ranges:
        if b >= total:
            raise ValueError(f"Chapter range out of bounds (document has {total} chapter(s)).")
        new_book = epub.EpubBook()
        new_book.set_identifier(str(uuid.uuid4()))
        title_meta = book.get_metadata("DC", "title")
        base_title = title_meta[0][0] if title_meta else "Untitled"
        new_book.set_title(f"{base_title} (part)")
        lang_meta = book.get_metadata("DC", "language")
        new_book.set_language(lang_meta[0][0] if lang_meta else "en")
        chapters = []
        for rid in spine_ids[a:b + 1]:
            src_item = doc_items.get(rid)
            if src_item is None:
                continue
            new_item = epub.EpubHtml(title=src_item.get_name(), file_name=src_item.get_name(), lang="en")
            new_item.set_content(src_item.get_content())
            new_book.add_item(new_item)
            chapters.append(new_item)
        for res in resource_items:
            try:
                new_book.add_item(res)
            except Exception:
                pass
        new_book.toc = chapters
        new_book.add_item(epub.EpubNcx())
        new_book.add_item(epub.EpubNav())
        new_book.spine = ["nav"] + chapters
        p = os.path.join(out_dir, f"tmp_{uuid.uuid4().hex}.epub")
        epub.write_epub(p, new_book)
        outs.append(p)
    return outs


def split_document(src_path, ext, ranges, target_ext, out_dir):
    """Splits a PDF or EPUB by page/chapter ranges, optionally converting each part."""
    target_ext = target_ext or ext
    if ext == "pdf":
        intermediate = _split_pdf(src_path, ranges, out_dir)
    elif ext == "epub":
        intermediate = split_epub(src_path, ranges, out_dir)
    else:
        raise ValueError("split_document only supports pdf/epub; use block-based splitting for other formats.")
    if target_ext == ext:
        return intermediate
    outputs = []
    for p in intermediate:
        blocks = extract_blocks(p, ext)
        newp = os.path.join(out_dir, f"tmp_{uuid.uuid4().hex}.{target_ext}")
        render_blocks(blocks, target_ext, newp)
        outputs.append(newp)
        try:
            os.remove(p)
        except OSError:
            pass
    return outputs


# ---- AI helpers -----------------------------------------------------------

SPLIT_SYSTEM_PROMPT = (
    "You are a precise document-structure analyst. You respond with strict JSON "
    "only \u2014 no prose, no markdown code fences, no explanations."
)
TOC_SYSTEM_PROMPT = (
    "You are an expert technical editor who writes exceptionally detailed, "
    "well-organized tables of contents in Markdown."
)
TOC_PROMPT_TEMPLATE = (
    "Create an extremely {depth} table of contents for the document below. "
    "Use nested Markdown headings/bullets that reflect the document's real structure, "
    "including sub-points and key ideas within each section where evident from the text. "
    "Respond with the table of contents only, in Markdown, and nothing else.\n\n"
    "DOCUMENT:\n{content}"
)


def resolve_api_key(provider, user_key):
    if user_key:
        return user_key.strip() or None
    return os.environ.get(ENV_KEYS.get(provider, ""), "").strip() or None


def call_claude(api_key, model, system, prompt, max_tokens=4000):
    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json",
            },
            json={
                "model": model, "max_tokens": max_tokens, "system": system,
                "messages": [{"role": "user", "content": prompt}],
            },
            timeout=120,
        )
    except requests.RequestException as e:
        raise RuntimeError(f"Could not reach the Claude API: {e}")
    if resp.status_code != 200:
        raise RuntimeError(f"Claude API error ({resp.status_code}): {resp.text[:300]}")
    data = resp.json()
    return "".join(b.get("text", "") for b in data.get("content", []) if b.get("type") == "text")


def call_gemini(api_key, model, system, prompt, max_tokens=4000):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
    payload = {
        "contents": [{"role": "user", "parts": [{"text": prompt}]}],
        "systemInstruction": {"parts": [{"text": system}]},
        "generationConfig": {"maxOutputTokens": max_tokens},
    }
    try:
        resp = requests.post(url, json=payload, timeout=120)
    except requests.RequestException as e:
        raise RuntimeError(f"Could not reach the Gemini API: {e}")
    if resp.status_code != 200:
        raise RuntimeError(f"Gemini API error ({resp.status_code}): {resp.text[:300]}")
    data = resp.json()
    candidates = data.get("candidates") or []
    if not candidates:
        raise RuntimeError("Gemini API returned no candidates (the request may have been blocked by safety filters).")
    parts = candidates[0].get("content", {}).get("parts", [])
    return "".join(p.get("text", "") for p in parts)


def call_ai(provider, key, model, system, prompt, max_tokens=4000):
    if provider == "claude":
        return call_claude(key, model, system, prompt, max_tokens)
    if provider == "gemini":
        return call_gemini(key, model, system, prompt, max_tokens)
    raise ValueError("Unknown AI provider.")


def truncate_text(s, n):
    if len(s) <= n:
        return s, False
    return s[:n] + "\n...[truncated for length]", True


def build_outline_units(blocks, ext):
    """Returns (outline_lines, total_units, unit_kind) where unit_kind is
    'page' (pdf), 'chapter' (epub), or 'block' (everything else)."""
    if ext == "pdf":
        pages = {}
        for b in blocks:
            pages.setdefault(b["page"], []).append(b)
        total = (max(pages.keys()) + 1) if pages else 0
        lines = []
        for i in range(total):
            pblocks = pages.get(i, [])
            heads = [x["text"] for x in pblocks if x["type"] == "heading"]
            snippet = next((x["text"] for x in pblocks if x["type"] == "para"), "")
            desc = "; ".join(heads[:2]) if heads else snippet[:100]
            lines.append(f"[{i}] {desc}")
        return lines, total, "page"
    if ext == "epub":
        chapters = {}
        for b in blocks:
            chapters.setdefault(b.get("chapter", 0), []).append(b)
        total = (max(chapters.keys()) + 1) if chapters else 0
        lines = []
        for i in range(total):
            cb = chapters.get(i, [])
            heads = [x["text"] for x in cb if x["type"] == "heading"]
            snippet = next((x["text"] for x in cb if x["type"] == "para"), "")
            desc = "; ".join(heads[:2]) if heads else snippet[:100]
            lines.append(f"[{i}] {desc}")
        return lines, total, "chapter"
    lines = []
    for i, b in enumerate(blocks):
        tag = f"H{b['level']}" if b["type"] == "heading" else "P"
        lines.append(f"[{i}] ({tag}) {b['text'][:100]}")
    return lines, len(blocks), "block"


def build_split_prompt(outline_text, unit_kind, total_units):
    return (
        f"Below is an outline of a document broken into {total_units} {unit_kind}s, numbered from 0.\n"
        "Identify natural chapter/section boundaries.\n"
        "Return ONLY valid JSON of this exact form: "
        '{"chapters": [{"title": "string", "start_unit": integer}, ...]}\n'
        f"Rules: start_unit values must be strictly increasing integers between 0 and {total_units - 1}. "
        "The first chapter must have start_unit = 0. Do not include any text outside the JSON.\n\n"
        f"OUTLINE:\n{outline_text}"
    )


def parse_split_response(raw, total_units):
    text = raw.strip()
    text = re.sub(r"^```(?:json)?\s*|\s*```$", "", text, flags=re.S).strip()
    try:
        data = json.loads(text)
    except Exception:
        m = re.search(r"\{.*\}", text, re.S)
        if not m:
            raise RuntimeError("The AI response was not valid JSON.")
        data = json.loads(m.group(0))
    chapters = data.get("chapters") if isinstance(data, dict) else None
    if not isinstance(chapters, list) or not chapters:
        raise RuntimeError("The AI response did not contain a usable chapters list.")
    cleaned = []
    seen = set()
    for c in chapters:
        if not isinstance(c, dict):
            continue
        try:
            su = int(c.get("start_unit"))
        except Exception:
            continue
        su = max(0, min(su, total_units - 1))
        if su in seen:
            continue
        seen.add(su)
        title = str(c.get("title") or "").strip()[:120] or f"Part {len(cleaned) + 1}"
        cleaned.append({"title": title, "start_unit": su})
    cleaned.sort(key=lambda x: x["start_unit"])
    if not cleaned:
        raise RuntimeError("The AI response did not contain any valid chapter boundaries.")
    if cleaned[0]["start_unit"] != 0:
        cleaned.insert(0, {"title": "Introduction", "start_unit": 0})
    return cleaned


def chapters_to_ranges(chapters, total_units):
    ranges = []
    for i, c in enumerate(chapters):
        start = c["start_unit"]
        end = (chapters[i + 1]["start_unit"] - 1) if i + 1 < len(chapters) else total_units - 1
        if end < start:
            end = start
        ranges.append((start, end))
    return ranges


def build_full_text(blocks, max_chars):
    parts = []
    for b in blocks:
        if b["type"] == "heading":
            parts.append(("#" * min(b["level"], 6)) + " " + b["text"])
        else:
            parts.append(b["text"])
    return truncate_text("\n".join(parts), max_chars)


# ---- File analysis ----------------------------------------------------------

def analyze_file(path, ext):
    meta = {}
    try:
        if ext == "pdf":
            if fitz is not None:
                d = fitz.open(path)
                n = len(d)
                d.close()
            elif PdfReader is not None:
                n = len(PdfReader(path).pages)
            else:
                n = None
            meta["unit_kind"] = "page"
            meta["units"] = n
        elif ext == "epub":
            if epub is not None:
                book = epub.read_epub(path)
                n = len([i for i in book.spine if i[0] != "nav"])
            else:
                n = None
            meta["unit_kind"] = "chapter"
            meta["units"] = n
        else:
            blocks = extract_blocks(path, ext)
            sections = sum(1 for b in blocks if b["type"] == "heading" and b["level"] == 1)
            meta["unit_kind"] = "section"
            meta["units"] = sections
            meta["blocks"] = len(blocks)
    except Exception as e:
        meta["analyze_error"] = str(e)
    return meta


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.route("/")
def index():
    return Response(PAGE_HTML, mimetype="text/html")


@app.route("/api/upload", methods=["POST"])
def api_upload():
    if "file" not in request.files:
        return jsonify(error="No file part in the request."), 400
    f = request.files["file"]
    if not f or f.filename == "":
        return jsonify(error="No file selected."), 400
    filename = secure_filename(f.filename)
    if not filename or "." not in filename:
        return jsonify(error="File must have a valid extension."), 400
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        return jsonify(error=f"Unsupported file type '.{ext}'. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}"), 400
    file_id = uuid.uuid4().hex
    stored_name = f"{file_id}.{ext}"
    dest = os.path.join(session_dir(), stored_name)
    try:
        f.save(dest)
    except Exception as e:
        return jsonify(error=f"Could not save upload: {e}"), 500
    meta = analyze_file(dest, ext)
    register_file(file_id, filename, ext, stored_name, meta)
    return jsonify({"file_id": file_id, "filename": filename, "ext": ext, "meta": meta})


@app.route("/api/convert", methods=["POST"])
def api_convert():
    data = request.get_json(force=True, silent=True) or {}
    file_id = data.get("file_id", "")
    target = (data.get("target_format") or "").lower()
    if not valid_id(file_id):
        return jsonify(error="Invalid file reference."), 400
    if target not in TARGET_FORMATS:
        return jsonify(error="Unsupported target format."), 400
    rec = get_file_record(file_id)
    if not rec:
        return jsonify(error="File not found or session expired. Please re-upload."), 404
    try:
        blocks = extract_blocks(rec["abs_path"], rec["ext"])
    except Exception as e:
        return jsonify(error=f"Could not read source document: {e}"), 400
    out_id = uuid.uuid4().hex
    stored_name = f"{out_id}.{target}"
    out_path = os.path.join(session_dir(), stored_name)
    try:
        render_blocks(blocks, target, out_path)
    except Exception as e:
        return jsonify(error=f"Conversion failed: {e}"), 500
    fname = f"{safe_stem(rec['filename'])}.{target}"
    register_output(out_id, fname, target, stored_name)
    return jsonify({"output_id": out_id, "filename": fname})


@app.route("/api/split", methods=["POST"])
def api_split():
    data = request.get_json(force=True, silent=True) or {}
    file_id = data.get("file_id", "")
    spec = data.get("spec", "")
    target = (data.get("target_format") or "").lower()
    if not valid_id(file_id):
        return jsonify(error="Invalid file reference."), 400
    rec = get_file_record(file_id)
    if not rec:
        return jsonify(error="File not found or session expired. Please re-upload."), 404
    if target and target not in TARGET_FORMATS:
        return jsonify(error="Unsupported target format."), 400
    target = target or rec["ext"]

    try:
        if rec["ext"] in ("pdf", "epub"):
            total = rec["meta"].get("units")
            if not total:
                return jsonify(error="Could not determine page/chapter count for this document."), 400
            ranges = parse_ranges(spec, total)
            tmp_paths = split_document(rec["abs_path"], rec["ext"], ranges, target_ext=target, out_dir=session_dir())
        else:
            blocks = extract_blocks(rec["abs_path"], rec["ext"])
            units = group_blocks_into_sections(blocks)
            if len(units) < 2:
                return jsonify(error="No top-level headings found to split by. Try 'Split smart by chapters (AI)' instead."), 400
            ranges = parse_ranges(spec, len(units))
            tmp_paths = []
            for a, b in ranges:
                sub = []
                for u in units[a:b + 1]:
                    sub.extend(u)
                p = os.path.join(session_dir(), f"tmp_{uuid.uuid4().hex}.{target}")
                render_blocks(sub, target, p)
                tmp_paths.append(p)
    except ValueError as e:
        return jsonify(error=str(e)), 400
    except Exception as e:
        return jsonify(error=f"Split failed: {e}"), 500

    results = []
    for i, p in enumerate(tmp_paths):
        oid = uuid.uuid4().hex
        stored_name = f"{oid}.{target}"
        final_path = os.path.join(session_dir(), stored_name)
        shutil.move(p, final_path)
        fname = f"{safe_stem(rec['filename'])}_part{i + 1}.{target}"
        register_output(oid, fname, target, stored_name)
        results.append({"output_id": oid, "filename": fname})
    return jsonify({"parts": results})


@app.route("/api/ai/split", methods=["POST"])
def api_ai_split():
    data = request.get_json(force=True, silent=True) or {}
    file_id = data.get("file_id", "")
    provider = data.get("provider", "")
    api_key = data.get("api_key") or None
    model = (data.get("model") or "").strip() or None
    target = (data.get("target_format") or "").lower()

    if not valid_id(file_id):
        return jsonify(error="Invalid file reference."), 400
    if provider not in ("claude", "gemini"):
        return jsonify(error="Unknown AI provider."), 400
    rec = get_file_record(file_id)
    if not rec:
        return jsonify(error="File not found or session expired. Please re-upload."), 404
    if target and target not in TARGET_FORMATS:
        return jsonify(error="Unsupported target format."), 400
    target = target or rec["ext"]

    key = resolve_api_key(provider, api_key)
    if not key:
        return jsonify(error=(
            f"No API key available for {provider}. Enter your own key, or set the "
            f"{ENV_KEYS[provider]} environment variable on the server."
        )), 400

    try:
        blocks = extract_blocks(rec["abs_path"], rec["ext"])
        lines, total_units, unit_kind = build_outline_units(blocks, rec["ext"])
        if total_units < 2:
            return jsonify(error="Document is too short to split into chapters."), 400
        outline_text, truncated = truncate_text("\n".join(lines), 15000)
        prompt = build_split_prompt(outline_text, unit_kind, total_units)
        raw = call_ai(provider, key, model or DEFAULT_MODELS[provider], SPLIT_SYSTEM_PROMPT, prompt)
        chapters = parse_split_response(raw, total_units)
        ranges = chapters_to_ranges(chapters, total_units)
        if rec["ext"] in ("pdf", "epub"):
            tmp_paths = split_document(rec["abs_path"], rec["ext"], ranges, target_ext=target, out_dir=session_dir())
        else:
            tmp_paths = []
            for a, b in ranges:
                sub = blocks[a:b + 1]
                p = os.path.join(session_dir(), f"tmp_{uuid.uuid4().hex}.{target}")
                render_blocks(sub, target, p)
                tmp_paths.append(p)
    except RuntimeError as e:
        return jsonify(error=str(e)), 502
    except ValueError as e:
        return jsonify(error=str(e)), 400
    except Exception as e:
        traceback.print_exc()
        return jsonify(error=f"Smart split failed: {e}"), 500

    results = []
    for i, p in enumerate(tmp_paths):
        oid = uuid.uuid4().hex
        stored_name = f"{oid}.{target}"
        final_path = os.path.join(session_dir(), stored_name)
        shutil.move(p, final_path)
        title = chapters[i]["title"] if i < len(chapters) else f"Part {i + 1}"
        fname = f"{sanitize_title(title)}.{target}"
        register_output(oid, fname, target, stored_name)
        results.append({"output_id": oid, "filename": fname, "title": title})
    return jsonify({"chapters": results, "truncated_outline": truncated})


@app.route("/api/ai/toc", methods=["POST"])
def api_ai_toc():
    data = request.get_json(force=True, silent=True) or {}
    file_id = data.get("file_id", "")
    provider = data.get("provider", "")
    api_key = data.get("api_key") or None
    model = (data.get("model") or "").strip() or None
    depth = data.get("depth") or "detailed, with subsections and key topics"

    if not valid_id(file_id):
        return jsonify(error="Invalid file reference."), 400
    if provider not in ("claude", "gemini"):
        return jsonify(error="Unknown AI provider."), 400
    rec = get_file_record(file_id)
    if not rec:
        return jsonify(error="File not found or session expired. Please re-upload."), 404

    key = resolve_api_key(provider, api_key)
    if not key:
        return jsonify(error=(
            f"No API key available for {provider}. Enter your own key, or set the "
            f"{ENV_KEYS[provider]} environment variable on the server."
        )), 400

    try:
        blocks = extract_blocks(rec["abs_path"], rec["ext"])
        content_text, truncated = build_full_text(blocks, 100000)
        prompt = TOC_PROMPT_TEMPLATE.format(depth=depth, content=content_text)
        raw = call_ai(provider, key, model or DEFAULT_MODELS[provider], TOC_SYSTEM_PROMPT, prompt, max_tokens=4000)
    except RuntimeError as e:
        return jsonify(error=str(e)), 502
    except Exception as e:
        traceback.print_exc()
        return jsonify(error=f"TOC generation failed: {e}"), 500

    out_id = uuid.uuid4().hex
    stored_name = f"{out_id}.md"
    out_path = os.path.join(session_dir(), stored_name)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(raw)
    fname = f"{safe_stem(rec['filename'])}_TOC.md"
    register_output(out_id, fname, "md", stored_name)
    return jsonify({"toc_text": raw, "output_id": out_id, "filename": fname, "truncated": truncated})


@app.route("/api/download/<kind>/<fid>")
def api_download(kind, fid):
    if not valid_id(fid):
        abort(404)
    if kind == "output":
        rec = get_output_record(fid)
    elif kind == "source":
        rec = get_file_record(fid)
    else:
        abort(404)
    if not rec:
        abort(404)
    return send_file(rec["abs_path"], as_attachment=True, download_name=rec["filename"])


@app.route("/api/download-all", methods=["POST"])
def api_download_all():
    data = request.get_json(force=True, silent=True) or {}
    ids = data.get("output_ids") or []
    buf = io.BytesIO()
    added = 0
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        used_names = set()
        for oid in ids:
            if not valid_id(oid):
                continue
            rec = get_output_record(oid)
            if not rec:
                continue
            name = rec["filename"]
            n = name
            i = 2
            while n in used_names:
                stem, dot, extn = name.rpartition(".")
                n = f"{stem}_{i}.{extn}" if dot else f"{name}_{i}"
                i += 1
            used_names.add(n)
            zf.write(rec["abs_path"], arcname=n)
            added += 1
    if added == 0:
        return jsonify(error="No valid files to package."), 400
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="documents.zip", mimetype="application/zip")


@app.errorhandler(413)
def too_large(e):
    return jsonify(error="File is too large (60 MB limit)."), 413


@app.errorhandler(500)
def server_error(e):
    return jsonify(error="Internal server error."), 500


# ---------------------------------------------------------------------------
# Frontend (single embedded HTML page: CSS + JS inline, no external assets)
# ---------------------------------------------------------------------------
PAGE_HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Document Utilities</title>
<style>
  :root{
    --bg:#E7EAEE; --bg-panel:#F4F5F3; --ink:#1B2430; --ink-soft:#4B5563;
    --line:#C7CDD6; --accent:#8A6A2F; --accent-ink:#5C4720; --accent-soft:#EFE6D2;
    --danger:#9A3B2E; --danger-bg:#F6E4E0; --radius:10px;
    --serif:'Iowan Old Style','Palatino Linotype',Palatino,Georgia,serif;
    --sans:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,Helvetica,Arial,sans-serif;
    --mono:ui-monospace,SFMono-Regular,Consolas,Menlo,monospace;
  }
  *{box-sizing:border-box}
  body{margin:0;background:var(--bg);color:var(--ink);font-family:var(--sans);
       -webkit-font-smoothing:antialiased;}
  .shell{max-width:880px;margin:0 auto;padding:28px 20px 60px}
  header.site{display:flex;align-items:baseline;justify-content:space-between;
       border-bottom:2px solid var(--ink);padding-bottom:14px;margin-bottom:22px}
  header.site h2{font-family:var(--serif);font-weight:600;font-size:1.7rem;margin:0;
       letter-spacing:.2px}
  header.site .kicker{font-family:var(--mono);font-size:.72rem;letter-spacing:.08em;
       text-transform:uppercase;color:var(--ink-soft)}
  .tab-bar{display:flex;gap:2px;margin-bottom:0;flex-wrap:wrap}
  .tab-btn{font-family:var(--mono);font-size:.78rem;letter-spacing:.03em;
       background:var(--bg);border:1px solid var(--line);border-bottom:none;
       color:var(--ink-soft);padding:10px 16px;cursor:pointer;border-radius:8px 8px 0 0;
       display:flex;align-items:center;gap:8px;transition:background .15s,color .15s}
  .tab-btn .n{display:inline-flex;align-items:center;justify-content:center;
       width:18px;height:18px;border-radius:50%;border:1px solid currentColor;font-size:.68rem}
  .tab-btn.active{background:var(--bg-panel);color:var(--ink);border-color:var(--ink);
       font-weight:600}
  .tab-btn:disabled{opacity:.4;cursor:not-allowed}
  .tab-btn:not(.active):not(:disabled):hover{color:var(--ink)}
  .panel{background:var(--bg-panel);border:1px solid var(--ink);border-radius:0 8px 8px 8px;
       padding:26px;min-height:220px}
  .tab-panel{display:none}
  .tab-panel.active{display:block}
  h3.panel-title{font-family:var(--serif);margin:0 0 6px;font-size:1.2rem}
  p.hint{color:var(--ink-soft);font-size:.88rem;margin:0 0 18px;line-height:1.5}
  .dropzone{border:2px dashed var(--line);border-radius:var(--radius);padding:38px 20px;
       text-align:center;cursor:pointer;transition:border-color .15s,background .15s;background:#FBFBF9}
  .dropzone.drag{border-color:var(--accent);background:var(--accent-soft)}
  .dropzone .big{font-family:var(--serif);font-size:1.05rem;margin-bottom:6px}
  .dropzone .small{font-family:var(--mono);font-size:.72rem;color:var(--ink-soft)}
  input[type=file]{display:none}
  .file-card{margin-top:16px;border:1px solid var(--line);border-radius:var(--radius);
       padding:14px 16px;display:flex;justify-content:space-between;align-items:center;
       background:#fff}
  .file-card .name{font-weight:600}
  .file-card .meta{font-family:var(--mono);font-size:.76rem;color:var(--ink-soft);margin-top:2px}
  select, input[type=text], input[type=password], textarea{
       font-family:var(--sans);font-size:.92rem;padding:9px 11px;border:1px solid var(--line);
       border-radius:7px;background:#fff;color:var(--ink);width:100%}
  select{font-family:var(--sans)}
  label.field{display:block;font-size:.78rem;color:var(--ink-soft);margin:14px 0 5px;
       font-family:var(--mono);text-transform:uppercase;letter-spacing:.04em}
  .op-list{display:flex;flex-direction:column;gap:8px;margin-top:6px}
  .op-item{border:1px solid var(--line);border-radius:var(--radius);padding:12px 14px;
       cursor:pointer;background:#fff;display:flex;align-items:center;gap:10px}
  .op-item input{margin:0}
  .op-item.selected{border-color:var(--accent);background:var(--accent-soft)}
  .op-item .lbl{font-weight:600;font-size:.94rem}
  .op-item .desc{font-size:.8rem;color:var(--ink-soft);margin-top:2px}
  .ctrl-block{display:none}
  .ctrl-block.active{display:block}
  .api-box{margin-top:18px;border:1px dashed var(--accent);border-radius:var(--radius);
       padding:14px 16px;background:var(--accent-soft)}
  .api-box .row{display:flex;gap:14px;align-items:center;flex-wrap:wrap;margin-bottom:8px}
  .api-box label.inline{display:flex;align-items:center;gap:7px;font-size:.86rem;cursor:pointer}
  .radio-row{display:flex;gap:16px;flex-wrap:wrap;margin-top:6px}
  .radio-row label{display:flex;align-items:center;gap:6px;font-size:.88rem;cursor:pointer}
  .btn{font-family:var(--mono);font-size:.82rem;letter-spacing:.03em;text-transform:uppercase;
       background:var(--ink);color:#fff;border:none;border-radius:7px;padding:11px 18px;
       cursor:pointer;transition:opacity .15s}
  .btn:hover{opacity:.85}
  .btn:disabled{opacity:.4;cursor:not-allowed}
  .btn.secondary{background:#fff;color:var(--ink);border:1px solid var(--ink)}
  .btn.accent{background:var(--accent);}
  .actions{margin-top:20px;display:flex;gap:10px;align-items:center;flex-wrap:wrap}
  .status{font-size:.85rem;margin-top:10px;min-height:1.2em}
  .status.err{color:var(--danger);background:var(--danger-bg);padding:8px 12px;border-radius:7px}
  .status.ok{color:#1E5C34}
  .status.busy{color:var(--ink-soft)}
  textarea{min-height:220px;font-family:var(--mono);font-size:.84rem;line-height:1.5;resize:vertical}
  .results-grid{display:flex;flex-direction:column;gap:10px;margin-top:14px}
  .ticket{display:flex;justify-content:space-between;align-items:center;background:#fff;
       border:1px solid var(--line);border-left:3px dashed var(--accent);border-radius:6px;
       padding:11px 14px}
  .ticket .fname{font-family:var(--mono);font-size:.86rem}
  .ticket .title{font-size:.78rem;color:var(--ink-soft);margin-top:2px}
  .ticket a.dl{font-family:var(--mono);font-size:.76rem;text-decoration:none;color:var(--accent-ink);
       border:1px solid var(--accent);border-radius:6px;padding:6px 10px;white-space:nowrap}
  .ticket a.dl:hover{background:var(--accent-soft)}
  .empty-note{color:var(--ink-soft);font-size:.88rem;font-style:italic}
  footer.site{margin-top:36px;padding-top:14px;border-top:1px solid var(--line);
       font-family:var(--mono);font-size:.72rem;color:var(--ink-soft);text-align:center}
  @media (max-width:600px){.panel{padding:18px}}
  :focus-visible{outline:2px solid var(--accent);outline-offset:2px}
</style>
</head>
<body>
<div class="shell">
  <header class="site">
    <h2>Document Utilities</h2>
    <span class="kicker">convert &middot; split &middot; outline</span>
  </header>

  <nav class="tab-bar">
    <button class="tab-btn active" data-tab="document"><span class="n">1</span>Document</button>
    <button class="tab-btn" data-tab="operation" disabled><span class="n">2</span>Operation</button>
    <button class="tab-btn" data-tab="controls" disabled><span class="n">3</span>Configure</button>
    <button class="tab-btn" data-tab="results" disabled><span class="n">4</span>Results</button>
  </nav>

  <main class="panel">

    <!-- TAB 1: DOCUMENT -->
    <section class="tab-panel active" id="tab-document">
      <h3 class="panel-title">Upload a document</h3>
      <p class="hint">Supports PDF, DOCX, EPUB, Markdown, HTML and TXT &mdash; up to 60&nbsp;MB.</p>
      <div class="dropzone" id="dropzone" tabindex="0" role="button" aria-label="Upload a file">
        <div class="big">Drop a file here, or click to choose one</div>
        <div class="small">.pdf &nbsp;.docx &nbsp;.epub &nbsp;.md &nbsp;.html &nbsp;.txt</div>
      </div>
      <input type="file" id="fileInput" accept=".pdf,.docx,.epub,.md,.markdown,.html,.htm,.txt">
      <div id="fileCard" style="display:none" class="file-card">
        <div>
          <div class="name" id="fileName"></div>
          <div class="meta" id="fileMeta"></div>
        </div>
        <button class="btn secondary" id="changeFileBtn" type="button">Change file</button>
      </div>
      <div class="status" id="uploadStatus"></div>
      <div class="actions">
        <button class="btn" id="toOperationBtn" disabled type="button">Choose operation &rarr;</button>
      </div>
    </section>

    <!-- TAB 2: OPERATION -->
    <section class="tab-panel" id="tab-operation">
      <h3 class="panel-title">Choose an operation</h3>
      <p class="hint">What would you like to do with this document?</p>
      <div class="op-list" id="opList">
        <label class="op-item" data-op="convert">
          <input type="radio" name="op" value="convert">
          <div><div class="lbl">Convert to another format</div>
          <div class="desc">Re-render headings and paragraphs into PDF, DOCX, EPUB, Markdown, HTML or TXT.</div></div>
        </label>
        <label class="op-item" data-op="split">
          <input type="radio" name="op" value="split">
          <div><div class="lbl">Split by page / range</div>
          <div class="desc">Cut a PDF by page ranges, an EPUB by chapter, or another format by top-level sections.</div></div>
        </label>
        <label class="op-item" data-op="ai-split">
          <input type="radio" name="op" value="ai-split">
          <div><div class="lbl">Split smart by chapters (AI)</div>
          <div class="desc">Let Claude or Gemini find natural chapter boundaries and split accordingly.</div></div>
        </label>
        <label class="op-item" data-op="ai-toc">
          <input type="radio" name="op" value="ai-toc">
          <div><div class="lbl">Make an extremely detailed TOC (AI)</div>
          <div class="desc">Ask Claude or Gemini to generate a deeply nested table of contents.</div></div>
        </label>
      </div>
      <div class="actions">
        <button class="btn secondary" data-tab="document" type="button">&larr; Back</button>
        <button class="btn" id="toControlsBtn" disabled type="button">Configure &rarr;</button>
      </div>
    </section>

    <!-- TAB 3: CONTROLS -->
    <section class="tab-panel" id="tab-controls">
      <h3 class="panel-title" id="controlsTitle">Configure</h3>
      <p class="hint" id="controlsHint"></p>

      <div class="ctrl-block" id="ctrl-convert">
        <label class="field">Target format</label>
        <select id="convertTarget">
          <option value="pdf">PDF</option>
          <option value="docx">DOCX (Word)</option>
          <option value="epub">EPUB</option>
          <option value="md">Markdown</option>
          <option value="html">HTML</option>
          <option value="txt">Plain text</option>
        </select>
        <div class="actions"><button class="btn accent" id="runConvertBtn" type="button">Convert</button></div>
        <div class="status" id="convertStatus"></div>
      </div>

      <div class="ctrl-block" id="ctrl-split">
        <label class="field" id="splitUnitLabel">Pages</label>
        <div class="radio-row">
          <label><input type="radio" name="splitMode" value="each" checked> Each unit as its own file</label>
          <label><input type="radio" name="splitMode" value="everyn"> Every N units</label>
          <label><input type="radio" name="splitMode" value="custom"> Custom ranges</label>
        </div>
        <div id="everyNWrap" style="display:none">
          <label class="field">N</label>
          <input type="text" id="everyN" value="1">
        </div>
        <div id="customWrap" style="display:none">
          <label class="field">Ranges (e.g. 1-3,5,7-9)</label>
          <input type="text" id="customRanges" placeholder="1-3,5,7-9">
        </div>
        <label class="field">Output format</label>
        <select id="splitTarget"></select>
        <div class="actions"><button class="btn accent" id="runSplitBtn" type="button">Split document</button></div>
        <div class="status" id="splitStatus"></div>
      </div>

      <div class="ctrl-block" id="ctrl-ai-split">
        <label class="field">Output format for each chapter</label>
        <select id="aiSplitTarget"></select>
        <div id="apiBoxSplit"></div>
        <div class="actions"><button class="btn accent" id="runAiSplitBtn" type="button">Analyze &amp; split</button></div>
        <div class="status" id="aiSplitStatus"></div>
      </div>

      <div class="ctrl-block" id="ctrl-ai-toc">
        <label class="field">Detail level</label>
        <select id="tocDepth">
          <option value="detailed, with subsections and key topics" selected>Detailed</option>
          <option value="extremely detailed and exhaustive, with sub-subsections and key topics per section">Extremely detailed</option>
          <option value="concise, top-level only">Concise (top level only)</option>
        </select>
        <div id="apiBoxToc"></div>
        <div class="actions"><button class="btn accent" id="runAiTocBtn" type="button">Generate TOC</button></div>
        <div class="status" id="aiTocStatus"></div>
      </div>

      <div class="actions">
        <button class="btn secondary" data-tab="operation" type="button">&larr; Back</button>
      </div>
    </section>

    <!-- TAB 4: RESULTS -->
    <section class="tab-panel" id="tab-results">
      <h3 class="panel-title">Results</h3>
      <div id="tocPreviewWrap" style="display:none">
        <label class="field">Preview</label>
        <textarea id="tocPreview" readonly></textarea>
        <div class="actions"><button class="btn secondary" id="copyTocBtn" type="button">Copy text</button></div>
      </div>
      <label class="field" style="margin-top:22px">Downloadable files</label>
      <div class="results-grid" id="resultsGrid">
        <div class="empty-note">Nothing produced yet.</div>
      </div>
      <div class="actions">
        <button class="btn" id="downloadAllBtn" disabled type="button">Download all (.zip)</button>
        <button class="btn secondary" data-tab="operation" type="button">Run another operation</button>
      </div>
    </section>

  </main>

  <footer class="site">&copy; <span id="year"></span> &mdash; Document Utilities. Local processing utility; API keys are never stored.</footer>
</div>

<template id="apiBoxTemplate">
  <div class="api-box">
    <div class="row">
      <label class="inline"><input type="checkbox" class="use-api" checked disabled> Use AI API</label>
      <select class="provider">
        <option value="claude">Claude (Anthropic)</option>
        <option value="gemini">Gemini (Google)</option>
      </select>
    </div>
    <div class="row">
      <label class="inline"><input type="checkbox" class="enter-key"> Enter my own API key</label>
    </div>
    <div class="key-wrap" style="display:none">
      <label class="field">API key</label>
      <input type="password" class="api-key" autocomplete="off" placeholder="sk-... / AIza...">
      <label class="field">Model (optional override)</label>
      <input type="text" class="model" placeholder="e.g. claude-sonnet-5">
    </div>
    <p class="hint" style="margin:8px 0 0">If left unchecked, the server's own API key (environment variable) is used, if configured.</p>
  </div>
</template>

<script>
(function(){
  "use strict";
  document.getElementById('year').textContent = new Date().getFullYear();

  var state = { file: null, op: null, outputs: [] };

  function $(sel, root){ return (root||document).querySelector(sel); }
  function $all(sel, root){ return Array.prototype.slice.call((root||document).querySelectorAll(sel)); }

  // ---- Tabs ----------------------------------------------------------
  function activateTab(name){
    $all('.tab-btn').forEach(function(b){ b.classList.toggle('active', b.dataset.tab===name); });
    $all('.tab-panel').forEach(function(p){ p.classList.toggle('active', p.id==='tab-'+name); });
  }
  $all('[data-tab]').forEach(function(el){
    el.addEventListener('click', function(){
      if (el.disabled) return;
      activateTab(el.dataset.tab);
    });
  });

  function unlockTab(name){
    var btn = $all('.tab-btn').filter(function(b){ return b.dataset.tab===name; })[0];
    if (btn) btn.disabled = false;
  }

  // ---- Status helpers --------------------------------------------------
  function setStatus(el, msg, kind){
    el.textContent = msg || '';
    el.className = 'status' + (kind ? ' ' + kind : '');
  }

  // ---- Upload ------------------------------------------------------------
  var dropzone = $('#dropzone'), fileInput = $('#fileInput');
  dropzone.addEventListener('click', function(){ fileInput.click(); });
  dropzone.addEventListener('keydown', function(e){ if(e.key==='Enter'||e.key===' ') fileInput.click(); });
  ['dragenter','dragover'].forEach(function(ev){
    dropzone.addEventListener(ev, function(e){ e.preventDefault(); dropzone.classList.add('drag'); });
  });
  ['dragleave','drop'].forEach(function(ev){
    dropzone.addEventListener(ev, function(e){ e.preventDefault(); dropzone.classList.remove('drag'); });
  });
  dropzone.addEventListener('drop', function(e){
    var f = e.dataTransfer.files[0];
    if (f) uploadFile(f);
  });
  fileInput.addEventListener('change', function(){
    if (fileInput.files[0]) uploadFile(fileInput.files[0]);
  });
  $('#changeFileBtn').addEventListener('click', function(){
    fileInput.value = '';
    $('#fileCard').style.display = 'none';
    dropzone.style.display = 'block';
    $('#toOperationBtn').disabled = true;
    state.file = null;
  });

  function unitLabel(meta){
    if (!meta || meta.units == null) return 'Could not determine document structure.';
    var kind = meta.unit_kind === 'page' ? 'page' : (meta.unit_kind === 'chapter' ? 'chapter' : 'top-level section');
    var n = meta.units;
    return n + ' ' + kind + (n===1 ? '' : 's') + ' detected.';
  }

  function uploadFile(f){
    var statusEl = $('#uploadStatus');
    setStatus(statusEl, 'Uploading & analyzing\u2026', 'busy');
    var fd = new FormData();
    fd.append('file', f);
    fetch('/api/upload', { method:'POST', body: fd })
      .then(function(r){ return r.json().then(function(j){ return {ok:r.ok, body:j}; }); })
      .then(function(res){
        if (!res.ok) { setStatus(statusEl, res.body.error || 'Upload failed.', 'err'); return; }
        state.file = res.body;
        dropzone.style.display = 'none';
        $('#fileCard').style.display = 'flex';
        $('#fileName').textContent = res.body.filename;
        $('#fileMeta').textContent = res.body.ext.toUpperCase() + '  \u00b7  ' + unitLabel(res.body.meta);
        setStatus(statusEl, '', '');
        $('#toOperationBtn').disabled = false;
        unlockTab('operation');
      })
      .catch(function(e){ setStatus(statusEl, 'Upload failed: ' + e, 'err'); });
  }

  $('#toOperationBtn').addEventListener('click', function(){ activateTab('operation'); });

  // ---- Operation selection --------------------------------------------
  $all('.op-item').forEach(function(item){
    item.addEventListener('click', function(){
      $all('.op-item').forEach(function(o){ o.classList.remove('selected'); });
      item.classList.add('selected');
      item.querySelector('input').checked = true;
      state.op = item.dataset.op;
      $('#toControlsBtn').disabled = false;
    });
  });
  $('#toControlsBtn').addEventListener('click', function(){
    renderControls();
    activateTab('controls');
  });

  // ---- Controls panel ---------------------------------------------------
  var TITLES = {
    convert: ['Convert to another format', 'Pick the target format.'],
    split: ['Split by page / range', 'Choose how to divide the document.'],
    'ai-split': ['Split smart by chapters', 'An AI model proposes chapter boundaries, then the document is split accordingly.'],
    'ai-toc': ['Make a detailed table of contents', 'An AI model reads the document and drafts a nested table of contents.']
  };

  function populateFormatSelect(sel, excludeCurrent){
    sel.innerHTML = '';
    var formats = [['pdf','PDF'],['docx','DOCX (Word)'],['epub','EPUB'],['md','Markdown'],['html','HTML'],['txt','Plain text']];
    formats.forEach(function(pair){
      var opt = document.createElement('option');
      opt.value = pair[0]; opt.textContent = pair[1];
      sel.appendChild(opt);
    });
    if (state.file && excludeCurrent) sel.value = state.file.ext === 'markdown' ? 'md' : state.file.ext;
  }

  function buildApiBox(container){
    var tpl = $('#apiBoxTemplate');
    container.innerHTML = '';
    var node = tpl.content.cloneNode(true);
    container.appendChild(node);
    var box = container.querySelector('.api-box');
    var enterKey = box.querySelector('.enter-key');
    var keyWrap = box.querySelector('.key-wrap');
    var providerSel = box.querySelector('.provider');
    var modelInput = box.querySelector('.model');
    enterKey.addEventListener('change', function(){
      keyWrap.style.display = enterKey.checked ? 'block' : 'none';
    });
    function updateModelPlaceholder(){
      modelInput.placeholder = providerSel.value === 'claude' ? 'e.g. claude-sonnet-5' : 'e.g. gemini-3.1-flash-lite';
    }
    providerSel.addEventListener('change', updateModelPlaceholder);
    updateModelPlaceholder();
    return box;
  }

  function readApiBox(box){
    return {
      provider: box.querySelector('.provider').value,
      api_key: box.querySelector('.enter-key').checked ? box.querySelector('.api-key').value.trim() : '',
      model: box.querySelector('.model').value.trim()
    };
  }

  var aiSplitBox = null, aiTocBox = null;

  function renderControls(){
    var op = state.op;
    $('#controlsTitle').textContent = TITLES[op][0];
    $('#controlsHint').textContent = TITLES[op][1];
    $all('.ctrl-block').forEach(function(b){ b.classList.remove('active'); });
    $('#ctrl-' + op).classList.add('active');

    if (op === 'convert') {
      populateFormatSelect($('#convertTarget'), false);
    }
    if (op === 'split') {
      var meta = state.file.meta || {};
      var kindLabel = meta.unit_kind === 'page' ? 'Pages' : (meta.unit_kind === 'chapter' ? 'Chapters' : 'Sections');
      $('#splitUnitLabel').textContent = kindLabel + ' \u2014 ' + unitLabel(meta);
      populateFormatSelect($('#splitTarget'), true);
    }
    if (op === 'ai-split') {
      populateFormatSelect($('#aiSplitTarget'), true);
      aiSplitBox = buildApiBox($('#apiBoxSplit'));
    }
    if (op === 'ai-toc') {
      aiTocBox = buildApiBox($('#apiBoxToc'));
    }
  }

  $all('input[name=splitMode]').forEach(function(r){
    r.addEventListener('change', function(){
      $('#everyNWrap').style.display = (r.value === 'everyn' && r.checked) ? 'block' : 'none';
      $('#customWrap').style.display = (r.value === 'custom' && r.checked) ? 'block' : 'none';
    });
  });

  function computeSplitSpec(){
    var mode = $all('input[name=splitMode]').filter(function(r){ return r.checked; })[0].value;
    var total = (state.file.meta || {}).units || 0;
    if (mode === 'each') {
      var arr = [];
      for (var i=1;i<=total;i++) arr.push(i);
      return arr.join(',');
    }
    if (mode === 'everyn') {
      var n = parseInt($('#everyN').value, 10) || 1;
      var parts = [];
      for (var s=1; s<=total; s+=n) {
        var e = Math.min(s+n-1, total);
        parts.push(s===e ? (''+s) : (s+'-'+e));
      }
      return parts.join(',');
    }
    return $('#customRanges').value.trim();
  }

  // ---- Results rendering -------------------------------------------------
  function addResults(items){
    items.forEach(function(it){ state.outputs.push(it); });
    renderResultsGrid();
    unlockTab('results');
  }

  function renderResultsGrid(){
    var grid = $('#resultsGrid');
    if (state.outputs.length === 0) {
      grid.innerHTML = '<div class="empty-note">Nothing produced yet.</div>';
      $('#downloadAllBtn').disabled = true;
      return;
    }
    grid.innerHTML = '';
    state.outputs.forEach(function(it){
      var row = document.createElement('div');
      row.className = 'ticket';
      var left = document.createElement('div');
      left.innerHTML = '<div class="fname">' + escapeHtml(it.filename) + '</div>' +
        (it.title ? '<div class="title">' + escapeHtml(it.title) + '</div>' : '');
      var a = document.createElement('a');
      a.className = 'dl';
      a.href = '/api/download/output/' + it.output_id;
      a.textContent = 'Download';
      row.appendChild(left);
      row.appendChild(a);
      grid.appendChild(row);
    });
    $('#downloadAllBtn').disabled = false;
  }

  function escapeHtml(s){
    var d = document.createElement('div');
    d.textContent = s;
    return d.innerHTML;
  }

  $('#downloadAllBtn').addEventListener('click', function(){
    var ids = state.outputs.map(function(o){ return o.output_id; });
    fetch('/api/download-all', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({ output_ids: ids })
    }).then(function(r){
      if (!r.ok) return r.json().then(function(j){ throw new Error(j.error || 'Download failed'); });
      return r.blob();
    }).then(function(blob){
      var url = URL.createObjectURL(blob);
      var a = document.createElement('a');
      a.href = url; a.download = 'documents.zip';
      document.body.appendChild(a); a.click(); a.remove();
      URL.revokeObjectURL(url);
    }).catch(function(e){ alert(e.message || e); });
  });

  // ---- Action buttons -----------------------------------------------------
  $('#runConvertBtn').addEventListener('click', function(){
    var btn = this, statusEl = $('#convertStatus');
    btn.disabled = true;
    setStatus(statusEl, 'Converting\u2026', 'busy');
    fetch('/api/convert', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({ file_id: state.file.file_id, target_format: $('#convertTarget').value })
    }).then(function(r){ return r.json().then(function(j){ return {ok:r.ok, body:j}; }); })
      .then(function(res){
        btn.disabled = false;
        if (!res.ok) { setStatus(statusEl, res.body.error, 'err'); return; }
        setStatus(statusEl, 'Done.', 'ok');
        addResults([res.body]);
        activateTab('results');
      }).catch(function(e){ btn.disabled=false; setStatus(statusEl, ''+e, 'err'); });
  });

  $('#runSplitBtn').addEventListener('click', function(){
    var btn = this, statusEl = $('#splitStatus');
    var spec;
    try { spec = computeSplitSpec(); } catch(e){ setStatus(statusEl, 'Invalid range.', 'err'); return; }
    if (!spec) { setStatus(statusEl, 'Please provide a range.', 'err'); return; }
    btn.disabled = true;
    setStatus(statusEl, 'Splitting\u2026', 'busy');
    fetch('/api/split', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({ file_id: state.file.file_id, spec: spec, target_format: $('#splitTarget').value })
    }).then(function(r){ return r.json().then(function(j){ return {ok:r.ok, body:j}; }); })
      .then(function(res){
        btn.disabled = false;
        if (!res.ok) { setStatus(statusEl, res.body.error, 'err'); return; }
        setStatus(statusEl, res.body.parts.length + ' file(s) created.', 'ok');
        addResults(res.body.parts);
        activateTab('results');
      }).catch(function(e){ btn.disabled=false; setStatus(statusEl, ''+e, 'err'); });
  });

  $('#runAiSplitBtn').addEventListener('click', function(){
    var btn = this, statusEl = $('#aiSplitStatus');
    var api = readApiBox(aiSplitBox);
    btn.disabled = true;
    setStatus(statusEl, 'Asking ' + api.provider + ' to analyze structure\u2026', 'busy');
    fetch('/api/ai/split', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({
        file_id: state.file.file_id, provider: api.provider, api_key: api.api_key,
        model: api.model, target_format: $('#aiSplitTarget').value
      })
    }).then(function(r){ return r.json().then(function(j){ return {ok:r.ok, body:j}; }); })
      .then(function(res){
        btn.disabled = false;
        if (!res.ok) { setStatus(statusEl, res.body.error, 'err'); return; }
        setStatus(statusEl, res.body.chapters.length + ' chapter(s) created.' +
          (res.body.truncated_outline ? ' (outline truncated for very long documents)' : ''), 'ok');
        addResults(res.body.chapters);
        activateTab('results');
      }).catch(function(e){ btn.disabled=false; setStatus(statusEl, ''+e, 'err'); });
  });

  $('#runAiTocBtn').addEventListener('click', function(){
    var btn = this, statusEl = $('#aiTocStatus');
    var api = readApiBox(aiTocBox);
    btn.disabled = true;
    setStatus(statusEl, 'Asking ' + api.provider + ' to draft a table of contents\u2026', 'busy');
    fetch('/api/ai/toc', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({
        file_id: state.file.file_id, provider: api.provider, api_key: api.api_key,
        model: api.model, depth: $('#tocDepth').value
      })
    }).then(function(r){ return r.json().then(function(j){ return {ok:r.ok, body:j}; }); })
      .then(function(res){
        btn.disabled = false;
        if (!res.ok) { setStatus(statusEl, res.body.error, 'err'); return; }
        setStatus(statusEl, 'Done.' + (res.body.truncated ? ' (document truncated for length)' : ''), 'ok');
        $('#tocPreviewWrap').style.display = 'block';
        $('#tocPreview').value = res.body.toc_text;
        addResults([{ output_id: res.body.output_id, filename: res.body.filename }]);
        activateTab('results');
      }).catch(function(e){ btn.disabled=false; setStatus(statusEl, ''+e, 'err'); });
  });

  $('#copyTocBtn').addEventListener('click', function(){
    var ta = $('#tocPreview');
    ta.select();
    navigator.clipboard && navigator.clipboard.writeText(ta.value).catch(function(){ document.execCommand('copy'); });
  });

})();
</script>
</body>
</html>
"""


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5034))
    debug = os.environ.get("FLASK_DEBUG", "0") == "1"
    app.run(host="0.0.0.0", port=port, debug=debug)

